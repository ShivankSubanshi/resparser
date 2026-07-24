"""
Core resume-parsing / job-matching logic.

Refactored out of resP.py so it can be reused by both the CLI script
and the FastAPI backend, with the job description and resume files
passed in as arguments instead of hardcoded / read from a fixed folder.
"""

import io
import json
import os
import time
from pathlib import Path

from dotenv import load_dotenv
from docx import Document
from groq import Groq
from pydantic import BaseModel
from pypdf import PdfReader

load_dotenv()

_api_key = os.getenv("GROQ_API_KEY")
if not _api_key:
    raise ValueError("GROQ_API_KEY is not set (check your .env file)")

client = Groq(api_key=_api_key)
MODEL = "openai/gpt-oss-120b"

# Delay between Groq calls to stay under rate limits. Two calls are made
# per resume (parse + score), so this adds ~10s of latency per resume.
LLM_CALL_DELAY_SECONDS = 5


# ---------------------------------------------------------------------------
# Schemas
# ---------------------------------------------------------------------------

class JobD(BaseModel):
    role: str
    required_skills: list[str]
    preferred_skills: list[str]
    minimum_experience: float | None
    education_requirements: list[str]
    responsibilities: list[str]


class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []


class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None
    total_experience_years: float | None = None
    skills: list[str] = []
    experiences: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []


class MatchResult(BaseModel):
    score: float
    details: dict


JOBD_SCHEMA = JobD.model_json_schema()
RESUME_SCHEMA = Resume.model_json_schema()


# ---------------------------------------------------------------------------
# File reading (accepts raw bytes now, not just disk paths, so uploaded
# files can be handled without ever touching disk)
# ---------------------------------------------------------------------------

def read_pdf(file) -> str:
    reader = PdfReader(file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text


def read_docx(file) -> str:
    document = Document(file)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


def read_resume_bytes(filename: str, content: bytes) -> str | None:
    """Read resume text from raw bytes, dispatching on file extension."""
    suffix = Path(filename).suffix.lower()
    stream = io.BytesIO(content)
    if suffix == ".pdf":
        return read_pdf(stream)
    elif suffix == ".docx":
        return read_docx(stream)
    return None


def read_resume_path(file_path: Path) -> str | None:
    """Kept for CLI/local-folder usage."""
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(str(file_path))
    elif file_path.suffix.lower() == ".docx":
        return read_docx(str(file_path))
    return None


# ---------------------------------------------------------------------------
# LLM calls
# ---------------------------------------------------------------------------

def extract_job_description(job_description_text: str) -> JobD:
    system_prompt = f"""
You are an expert HR assistant.

Your job is to analyze job descriptions and extract
structured information from them.

Return ONLY valid JSON matching this schema:

{JOBD_SCHEMA}
IMPORTANT:
Do NOT return the schema itself.
Do NOT return fields like "properties", "title" or "type".
Fill the schema with actual information extracted from the job description.

If minimum experience is not mentioned, return null.
If information for a list is missing, return an empty list.
Do not invent information.
"""
    user_prompt = f"""
Analyze the following job description:

{job_description_text}
"""
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]
    response = client.chat.completions.create(
        model=MODEL,
        messages=messages,
        response_format={"type": "json_object"},
    )
    data = json.loads(response.choices[0].message.content)
    return JobD(**data)


def parse_resume(resume_text: str) -> Resume:
    system_prompt = f"""
You are an expert resume parser.

Extract information from the resume based on its meaning,
not only based on exact section headings.

Different resumes may use different headings.

For example:
- Experience
- Professional Experience
- Work History
- Employment
- Internships

These may all contain relevant experience.

Skills may also appear in the skills section, work experience,
internships or projects.

Return ONLY valid JSON matching this schema:

{RESUME_SCHEMA}

Important rules:

1. Do not invent information.
2. If a value is not available, return null.
3. If a list has no information, return an empty list.
4. Include internships inside experiences.
5. Extract skills mentioned across the entire resume.
"""
    user_prompt = f"""
Parse the following resume:

{resume_text}
"""
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]
    response = client.chat.completions.create(
        model=MODEL,
        messages=messages,
        response_format={"type": "json_object"},
    )
    data = json.loads(response.choices[0].message.content)
    return Resume(**data)


def final_score(job: JobD, resume: Resume) -> MatchResult:
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
    You are an HR recruiter.

    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    Return JSON matching this schema:

    {match_schema}

    Give me:

    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. Overall match percentage from 0 to 100
    6. A short final verdict

    Keep the response concise and easy to read.
    """
    messages = [{"role": "user", "content": prompt}]
    response = client.chat.completions.create(
        model=MODEL,
        messages=messages,
        response_format={"type": "json_object"},
    )
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)


# ---------------------------------------------------------------------------
# Orchestration used by the API (and reusable by the CLI script)
# ---------------------------------------------------------------------------

def analyze_resumes(
    job_description_text: str,
    files: list[tuple[str, bytes]],
) -> tuple[JobD, list[dict]]:
    """
    files: list of (filename, raw_bytes) tuples.
    Returns (parsed job description, list of per-candidate result dicts),
    sorted by score descending. Candidates that fail to parse/score are
    included with score=None and an "error" key in details, sorted last.
    """
    job = extract_job_description(job_description_text)

    results = []
    for filename, content in files:
        resume_text = read_resume_bytes(filename, content)
        if not resume_text:
            results.append({
                "name": filename,
                "score": None,
                "details": {"error": "Unsupported file type or empty file"},
            })
            continue
        try:
            parsed_resume = parse_resume(resume_text)
            time.sleep(LLM_CALL_DELAY_SECONDS)
            match = final_score(job, parsed_resume)
            time.sleep(LLM_CALL_DELAY_SECONDS)
            results.append({
                "name": parsed_resume.name or filename,
                "score": match.score,
                "details": match.details,
            })
        except Exception as exc:
            results.append({
                "name": filename,
                "score": None,
                "details": {"error": str(exc)},
            })

    results.sort(
        key=lambda candidate: candidate["score"] if candidate["score"] is not None else -1,
        reverse=True,
    )
    return job, results
